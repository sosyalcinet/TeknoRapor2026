import streamlit as st
import google.generativeai as genai
from docx import Document
from docx.shared import Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH
from fpdf import FPDF
from io import BytesIO
from datetime import datetime

# --- 1. RESMİ ŞARTNAME VERİ MATRİSİ (2026 GÜNCEL) ---
# Şartnamelerden çekilen dinamik veriler [cite: 1308-1315, 1617-1642, 2047-2061]
TEKNOFEST_MATRIS = {
    "İlkokul": {
        "Yarisma": "2026 İNSANLIK YARARINA TEKNOLOJİLER YARIŞMASI İLKOKUL SEVİYESİ",
        "Temalar": {
            "Doğa, Çevre ve Sürdürülebilirlik": ["Atık Yönetimi ve Geri Dönüşüm", "Yeşil Teknolojiler ve Yenilenebilir Enerji", "Akıllı Şehirler", "Doğal Yaşam", "Afet Teknolojileri"],
            "Astronomi, Uzay Bilimleri ve Havacılık": ["Gezegenimiz ve Evren", "Gözlemler ve Teleskoplar", "Uydular ve Roketler", "Uzayda Yaşam"],
            "Sağlıklı Yaşam": ["Fiziksel ve Zihinsel Sağlık", "Besin (Gıda) Teknolojileri", "Sürdürülebilir Tarım"]
        },
        "Hedef_Kitle": ["İlkokul Öğrencileri", "Veliler", "Öğretmenler", "Okul Çalışanları"]
    },
    "Ortaokul": {
        "Yarisma": "2026 İNSANLIK YARARINA TEKNOLOJİLER YARIŞMASI ORTAOKUL SEVİYESİ",
        "Temalar": {
            "Astronomi ve Uzay Teknolojileri": ["Uzay Araçları ve Keşif", "Gezegenler", "Gözlem Teknolojileri", "Evreni Keşfetmek"],
            "Doğa Bilimleri ve Çevresel Farkındalık": ["Akıllı Şehirler", "Ekosistemler", "Afetler ve Güvenli Yaşam", "Enerji Kaynakları", "Atık Yönetimi"],
            "Sağlık ve İyi Yaşam Teknolojileri": ["Beslenme ve Gıda", "Hareketli Yaşam", "Günlük Sağlık", "Zihinsel Sağlık", "Engelsiz Yaşam"],
            "Eğitim Teknolojileri": ["Dijital Araçlar", "Oyunlaştırma", "Dijital Güvenlik", "Öğrenmeyi Kolaylaştıran Çözümler"]
        },
        "Hedef_Kitle": ["Ortaokul Öğrencileri", "Bedensel Engelliler", "Yaşlılar", "Afetzedeler"]
    },
    "Lise": {
        "Yarisma": "2026 İNSANLIK YARARINA TEKNOLOJİLER YARIŞMASI LİSE SEVİYESİ",
        "Temalar": {
            "Akıllı Teknolojiler ve Sistem Tasarımı": ["Ulaşım ve Mobilite", "Şehir ve Kentsel Sistem", "Afet ve Acil Durum"],
            "Sağlık ve İyi Yaşam Teknolojileri": ["Hasta Odaklı", "Sağlık Çalışanlarına Yönelik", "İleri Araştırma", "Güvenli Yaşam"],
            "Eğitim, Kültür ve Dijital Deneyim Teknolojileri": ["Dijital Eğitim", "Etkileşimli Öğrenme", "Kültürel Miras"]
        },
        "Hedef_Kitle": ["Lise Öğrencileri", "Kronik Hastalar", "Profesyonel Kurtarma Ekipleri", "Yerel Yönetimler"]
    }
}

# GÜVENLİK VE ANAHTAR (Yeni Key Kullanımı)
try:
    # Kullanıcıdan gelen yeni anahtar: AIzaSyBLG5jhEOO44BU_BfIVVSz7L64AAIi7qBs
    GEMINI_API_KEY = st.secrets.get("GEMINI_API_KEY", "AIzaSyBLG5jhEOO44BU_BfIVVSz7L64AAIi7qBs")
except:
    st.error("Kasa ayarları bulunamadı!")
    st.stop()

st.set_page_config(page_title="TeknoRapor V20 | Derepazarı", layout="centered", page_icon="🤖")

# --- KARAKTER FİLTRESİ ---
def format_temizle(text):
    harita = {'İ': 'I', 'ı': 'i', 'Ş': 'S', 'ş': 's', 'Ğ': 'G', 'ğ': 'g', 'Ç': 'C', 'ç': 'c', 'Ö': 'O', 'ö': 'o', 'Ü': 'U', 'ü': 'u'}
    for k, v in harita.items(): text = text.replace(k, v)
    return text.encode('latin-1', 'replace').decode('latin-1')

# --- WORD FORMATLAMA (Arial 12pt, 1.15 Aralık) ---
def create_word_official(text, info):
    doc = Document()
    # 1. SAYFA: KAPAK [cite: 1231-1234, 1559, 1990]
    h_yarisma = doc.add_heading(info['y_adi'], 0)
    h_yarisma.alignment = WD_ALIGN_PARAGRAPH.CENTER
    doc.add_heading("ÖN DEĞERLENDİRME RAPORU", 1).alignment = WD_ALIGN_PARAGRAPH.CENTER
    p = doc.add_paragraph()
    p.add_run(f"\nPROJE ADI: {info['p_adi']}\nTAKIM ADI: {info['takim']}\nBAŞVURU ID: {info['b_id']}\nTAKIM ID: {info['t_id']}").bold = True
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    doc.add_page_break()
    # 2. SAYFA: İÇİNDEKİLER
    doc.add_heading("İÇİNDEKİLER", 1)
    doc.add_paragraph("1. PROJE ÖZETİ\n2. PROBLEMİN TANIMI VE ÇÖZÜM ÖNERİSİ\n3. ÖZGÜNLÜK VE UYGULANABİLİRLİK\n4. YÖNTEM VE SÜREÇ\n5. PROJE TAKIMI\n6. KAYNAKLAR")
    doc.add_page_break()
    # 3. SAYFA: İÇERİK
    content = doc.add_paragraph(text.replace("**", "").replace("##", ""))
    style = doc.styles['Normal']
    style.font.name = 'Arial'; style.font.size = Pt(12)
    content.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    content.paragraph_format.line_spacing = 1.15
    bio = BytesIO(); doc.save(bio); return bio.getvalue()

# --- PDF HATASINI ÇÖZEN YENİ FONKSİYON ---
def create_pdf_fixed(text, title):
    pdf = FPDF()
    pdf.add_page()
    pdf.set_font("Arial", 'B', 14)
    pdf.multi_cell(0, 10, txt=format_temizle(title), align='C')
    pdf.ln(5)
    pdf.set_font("Arial", size=11)
    pdf.multi_cell(0, 7.5, txt=format_temizle(text))
    # AttributeError: 'FPDF' object has no attribute 'encode' hatasını çözen yöntem
    return pdf.output()

# --- ARAYÜZ ---
st.markdown("<h1 style='text-align: center; color: #1E3A8A;'>Teknofest Resmi ÖDR Robotu V20</h1>", unsafe_allow_html=True)

st.write("**Rapor Kaç Sayfa Olsun?**")
hedef_sayfa = st.radio("Sayfa", options=[1, 2, 3, 4, 5, 6], index=2, horizontal=True, label_visibility="collapsed")

st.markdown("### 🛠️ PROJE GİRİŞİ")

with st.expander("👥 Takım ve Kayıt Bilgileri", expanded=True):
    c1, c2, c3 = st.columns(3)
    t_adi = c1.text_input("Takım Adı", placeholder="Takım İsmi")
    b_id = c2.text_input("Başvuru ID", placeholder="ID No")
    t_id = c3.text_input("Takım ID", placeholder="Takım No")

with st.expander("🏷️ Seviye ve Kategori Seçimi", expanded=True):
    seviye = st.selectbox("Eğitim Seviyesi", list(TEKNOFEST_MATRIS.keys()))
    col_a, col_b = st.columns(2)
    ana_t = col_a.selectbox("Ana Tema", list(TEKNOFEST_MATRIS[seviye]["Temalar"].keys()))
    alt_t = col_a.selectbox("Alt Tema Seçin", TEKNOFEST_MATRIS[seviye]["Temalar"][ana_t])
    h_kitle = col_b.selectbox("Hedef Kitle Seçin", TEKNOFEST_MATRIS[seviye]["Hedef_Kitle"])
    danisman = col_b.text_input("Danışman Adı", placeholder="AD SOYAD")

# PROJE ADI VE ÖZETİ - İSTEDİĞİNİZ GİBİ BİR ARADA
with st.expander("📝 Proje Detayı ve Yazım Ayarları", expanded=True):
    p_adi_input = st.text_input("Proje Adı", placeholder="Akıllı Projenizin İsmi")
    aciklama = st.text_area("Projenizin Ana Fikrini Yazın (Proje Özeti)", height=150)
    yazim_modu = st.selectbox("Yazım Karakteri", ["Standart AI Dedektör (İnsan Gibi Yaz)", "Akademik/Resmi", "Süper AI"])

    if st.button("🚀 Şartnameye Uygun Raporu Hazırla", use_container_width=True, type="primary"):
        if not aciklama or not p_adi_input:
            st.warning("Lütfen Proje Adı ve Açıklamasını doldurun.")
        else:
            with st.status(f"🛠️ {seviye} Seviyesi Raporu İşleniyor...", expanded=True) as status:
                try:
                    # 404 HATASINI ÇÖZEN EN STABİL MODEL ÇAĞRISI
                    genai.configure(api_key=GEMINI_API_KEY)
                    model = genai.GenerativeModel('gemini-1.5-flash')
                    
                    info_dict = {
                        "y_adi": TEKNOFEST_MATRIS[seviye]["Yarisma"],
                        "takim": t_adi if t_adi else "________________",
                        "b_id": b_id if b_id else "________________",
                        "t_id": t_id if t_id else "________________",
                        "p_adi": p_adi_input,
                        "kategori": f"{ana_t} / {alt_t}"
                    }

                    prompt = f"""
                    Sen profesyonel bir Teknofest danışmanısın. {seviye} seviyesi için resmi ÖDR yaz.
                    MOD: {yazim_modu} (İnsansı bir dil kullan).
                    Hedef: {hedef_sayfa} sayfa. Tema: {ana_t}/{alt_t}. Hedef Kitle: {h_kitle}.
                    ---
                    BÖLÜMLER (Puanlama Odaklı):
                    1. PROJE ÖZETİ (20p): {ana_t} uyumu.
                    2. PROBLEMİN TANIMI VE ÇÖZÜM ÖNERİSİ (35p).
                    3. ÖZGÜNLÜK VE UYGULANABİLİRLİK (24p).
                    4. YÖNTEM VE SÜREÇ (12p).
                    5. TAKIM VE KAYNAKLAR (6p).
                    ---
                    İçerik Temeli: {aciklama}
                    """
                    response = model.generate_content(prompt)
                    st.session_state.rapor = response.text
                    st.session_state.info = info_dict
                    st.session_state.hazir = True
                    status.update(label="✅ Rapor Başarıyla Hazırlandı!", state="complete")
                except Exception as e:
                    st.error(f"Sistem Pürüzü: {str(e)}")

if "hazir" in st.session_state:
    st.markdown("---")
    st.markdown(f"<div style='background:white; padding:30px; color:black; border:1px solid #ddd; border-radius:10px; font-family:Arial; text-align:justify;'>{st.session_state.rapor.replace('\n', '<br>')}</div>", unsafe_allow_html=True)
    
    col_d1, col_d2 = st.columns(2)
    with col_d1:
        st.download_button("📥 PDF İndir", create_pdf_fixed(st.session_state.rapor, st.session_state.info['p_adi']), f"{st.session_state.info['p_adi']}.pdf", use_container_width=True)
    with col_d2:
        st.download_button("📥 Word İndir (Kapaklı)", create_word_official(st.session_state.rapor, st.session_state.info), f"{st.session_state.info['p_adi']}.docx", use_container_width=True)

st.markdown(f"<p style='text-align: center; color: gray;'>Derepazarı İlçe MEM Arge Birimi © 2026<br><b>Hazırlayan: Hüsamettin KAYMAKÇI</b></p>", unsafe_allow_html=True)
